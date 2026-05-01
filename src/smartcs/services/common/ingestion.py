"""文档摄入管道

6 阶段流水线：Parse → Clean → Chunk → Embed → Dual-Write → Publish
将原始文档转换为可检索的知识块，写入 ES + Milvus 并发布 Kafka 事件。
"""

from __future__ import annotations

import asyncio
import hashlib
import logging
import re
import time
from datetime import datetime
from typing import TYPE_CHECKING, Any

from markdown_it import MarkdownIt

if TYPE_CHECKING:
    from sqlalchemy.ext.asyncio import AsyncSession

    from smartcs.services.common.embedding import EmbeddingProvider

from smartcs.shared.exceptions import DocumentFormatError
from smartcs.shared.models import DocumentMetadata
from smartcs.shared.orm_models import (
    KbChunk,
    KbDocStatus,
    KbDocument,
    KbEmbedStatus,
    KbIngestionLog,
    KbIngestionStage,
    KbIngestionStatus,
    KbSourceType,
)

logger = logging.getLogger(__name__)


# ── 常量 ──

# 中文句子结尾（优先断点）
_SENTENCE_ENDINGS = "。！？；\n"
# 中文短语结尾（次优先断点）
_PHRASE_ENDINGS = "，、：\u201c\u201d）】》"


# ══════════════════════════════════════════════════════════════
# 1. Parse 阶段
# ══════════════════════════════════════════════════════════════


def parse_markdown(content: str) -> str:
    """从 Markdown 文本中提取纯文本

    使用 markdown-it-py 解析 token 树，提取 inline/text 节点内容。
    """
    md = MarkdownIt()
    tokens = md.parse(content)
    parts: list[str] = []
    for token in tokens:
        if token.type == "inline":
            for child in token.children or []:
                if child.type == "text":
                    parts.append(child.content)
        elif (
            token.type in ("heading_open", "paragraph_open", "bullet_list_open", "ordered_list_open")
            and parts
            and parts[-1] != "\n"
        ):
            # 在块级元素之间插入换行
            parts.append("\n")
    return "\n".join(line.strip() for line in "".join(parts).split("\n") if line.strip())


def parse_html(content: str) -> str:
    """从 HTML 文本中提取纯文本

    移除 script/style/nav/footer/header 标签后提取文本。
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")
    for tag in soup.find_all(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def parse_pdf(file_path: str) -> str:
    """从 PDF 文件中提取文本

    使用 pymupdf (fitz) 逐页提取。
    """
    import fitz

    doc = fitz.open(file_path)
    parts: list[str] = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts)


def parse_docx(file_path: str) -> str:
    """从 DOCX 文件中提取文本

    提取段落和表格内容。
    """
    from docx import Document

    doc = Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_xlsx(file_path: str) -> str:
    """从 XLSX 文件中提取文本

    每行格式为 "header: value | ..."。
    """
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h) if h is not None else "" for h in rows[0]]
        for row in rows[1:]:
            cells: list[str] = []
            for header, value in zip(headers, row, strict=False):
                if value is not None:
                    cells.append(f"{header}: {value}")
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def parse_text_content(content: str) -> str:
    """纯文本直通，仅做 strip"""
    return content.strip()


# ── Parse 调度 ──

_PARSE_DISPATCH: dict[str, Any] = {
    KbSourceType.MARKDOWN: parse_markdown,
    KbSourceType.HTML: parse_html,
    KbSourceType.PDF: parse_pdf,
    KbSourceType.DOCX: parse_docx,
    KbSourceType.XLSX: parse_xlsx,
    KbSourceType.TXT: parse_text_content,
}


def _parse(source_type: KbSourceType, file_path: str) -> str:
    """根据来源类型选择对应的解析器

    Args:
        source_type: 文档来源类型
        file_path: 文件路径（MARKDOWN 类型时为文本内容）

    Returns:
        解析后的纯文本
    """
    parser = _PARSE_DISPATCH.get(source_type)
    if parser is None:
        raise DocumentFormatError(f"不支持的文档格式: {source_type.value}")

    if source_type == KbSourceType.MARKDOWN:
        # MARKDOWN 类型: file_path 参数实际传入的是文本内容
        return parser(file_path)
    if source_type in (KbSourceType.HTML, KbSourceType.TXT):
        # HTML/TXT: file_path 参数也是文本内容
        return parser(file_path)
    # 文件类型: file_path 是磁盘路径
    return parser(file_path)


# ══════════════════════════════════════════════════════════════
# 2. Clean 阶段
# ══════════════════════════════════════════════════════════════

# 页眉页脚正则
_RE_PAGE_HEADER_FOOTER = re.compile(
    r"(第\s*\d+\s*页\s*/?\s*共\s*\d+\s*页|Page\s+\d+\s+of\s+\d+)",
    re.IGNORECASE,
)

# 控制字符（保留 \n \t \r）
_RE_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

# 连续空格
_RE_MULTI_SPACES = re.compile(r" {2,}")

# 3+ 换行
_RE_MULTI_NEWLINES = re.compile(r"\n{3,}")


def clean_text(raw: str) -> str:
    """文本清洗

    - 移除页眉页脚（第X页/共Y页, Page X of Y）
    - 移除控制字符（保留 \\n \\t）
    - 移除连续空格
    - 折叠 3+ 换行为 2
    - 段落级 MD5 去重
    """
    # 1. 移除页眉页脚
    text = _RE_PAGE_HEADER_FOOTER.sub("", raw)
    # 2. 移除控制字符
    text = _RE_CONTROL_CHARS.sub("", text)
    # 3. 移除连续空格
    text = _RE_MULTI_SPACES.sub(" ", text)
    # 4. 折叠 3+ 换行为 2
    text = _RE_MULTI_NEWLINES.sub("\n\n", text)

    # 5. 段落级 MD5 去重
    paragraphs = text.split("\n")
    seen_hashes: set[str] = set()
    deduped: list[str] = []
    for para in paragraphs:
        stripped = para.strip()
        if not stripped:
            deduped.append(para)
            continue
        h = hashlib.md5(stripped.encode()).hexdigest()  # 去重用途，非安全哈希
        if h not in seen_hashes:
            seen_hashes.add(h)
            deduped.append(para)

    return "\n".join(deduped).strip()


# ══════════════════════════════════════════════════════════════
# 3. Chunk 阶段
# ══════════════════════════════════════════════════════════════


def _find_break_point(text: str, start: int, chunk_size: int) -> int:
    """在目标位置附近搜索最佳断点

    优先在中文句号等句末标点处断开，其次在短语标点处断开。
    搜索范围为目标位置前后 200 字符。
    """
    target = start + chunk_size
    if target >= len(text):
        return len(text)

    # 搜索窗口: [target - 200, min(target + 200, len)]
    search_start = max(start, target - 200)
    search_end = min(len(text), target + 200)

    # 优先级 1: 句子结尾
    best_pos = -1
    for i in range(search_end - 1, search_start - 1, -1):
        if text[i] in _SENTENCE_ENDINGS:
            best_pos = i + 1  # 断点在标点之后
            break

    if best_pos != -1 and abs(best_pos - target) <= 200:
        return best_pos

    # 优先级 2: 短语结尾
    for i in range(search_end - 1, search_start - 1, -1):
        if text[i] in _PHRASE_ENDINGS:
            best_pos = i + 1
            break

    if best_pos != -1 and abs(best_pos - target) <= 200:
        return best_pos

    # 优先级 3: 空格
    for i in range(search_end - 1, search_start - 1, -1):
        if text[i] == " ":
            best_pos = i + 1
            break

    if best_pos != -1 and abs(best_pos - target) <= 200:
        return best_pos

    # 最终回退: 强制在 target 处断开
    return target


def chunk_text(text: str, chunk_size: int = 1500, overlap: int = 200) -> list[str]:
    """递归字符分割器

    优先在中文句号等句末标点处断开，其次在短语标点处断开。
    搜索断点范围为 200 字符，支持重叠回退。

    Args:
        text: 待分块文本
        chunk_size: 目标块大小（字符数）
        overlap: 重叠字符数

    Returns:
        分块文本列表
    """
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0

    while start < len(text):
        break_pos = _find_break_point(text, start, chunk_size)

        # 如果找不到有效断点（剩余文本不足一个块），取到末尾
        if break_pos <= start:
            break_pos = len(text)

        chunk = text[start:break_pos]
        if chunk.strip():
            chunks.append(chunk)

        # 计算下一块起始位置（回退 overlap 字符实现重叠）
        if break_pos >= len(text):
            break
        start = break_pos - overlap
        # 避免回退到当前块之前
        if start <= 0 and len(chunks) == 1:
            start = break_pos

    return chunks


# ══════════════════════════════════════════════════════════════
# 4. Embed 阶段
# ══════════════════════════════════════════════════════════════


async def embed_chunks(
    chunks: list[str],
    provider: EmbeddingProvider,
    batch_size: int = 128,
) -> list[list[float]]:
    """批量嵌入文本块

    按 batch_size 分批调用 EmbeddingProvider.embed()。

    Args:
        chunks: 待嵌入文本块列表
        provider: 嵌入服务提供者
        batch_size: 每批处理数量

    Returns:
        嵌入向量列表，与 chunks 顺序对应
    """
    all_embeddings: list[list[float]] = []
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        embeddings = await provider.embed(batch)
        all_embeddings.extend(embeddings)
    return all_embeddings


# ══════════════════════════════════════════════════════════════
# 5. Dual-Write 阶段
# ══════════════════════════════════════════════════════════════


async def write_to_es(
    chunks: list[dict],
    es_client: Any,
    index_name: str = "smartcs_knowledge",
) -> int:
    """将文本块写入 Elasticsearch

    Args:
        chunks: 待写入块列表，每个包含 chunk_id, doc_id, content, metadata 等字段
        es_client: Elasticsearch 异步客户端
        index_name: ES 索引名

    Returns:
        成功写入的块数
    """
    success_count = 0
    for chunk in chunks:
        try:
            await es_client.index(
                index=index_name,
                id=chunk["chunk_id"],
                document={
                    "chunk_id": chunk["chunk_id"],
                    "doc_id": chunk["doc_id"],
                    "content": chunk["content"],
                    "category": chunk.get("category", ""),
                    "doc_type": chunk.get("doc_type", ""),
                    "keywords": chunk.get("keywords", []),
                    "card_type": chunk.get("card_type", ""),
                    "customer_tier": chunk.get("customer_tier", ""),
                    "effective_date": chunk.get("effective_date"),
                    "expiry_date": chunk.get("expiry_date"),
                    "security_level": chunk.get("security_level", "internal"),
                    "version": chunk.get("version", "1.0"),
                },
            )
            success_count += 1
        except Exception:
            logger.exception("ES 写入失败: chunk_id=%s", chunk["chunk_id"])
    return success_count


async def write_to_milvus(
    chunks: list[dict],
    collection: Any,
) -> int:
    """将文本块写入 Milvus

    字段: chunk_id, doc_id, content, embedding, category, doc_type, keywords,
    card_type, customer_tier, effective_date(int64 epoch), expiry_date(int64 epoch)

    Args:
        chunks: 待写入块列表，每个包含 embedding 和 metadata 字段
        collection: Milvus Collection 对象

    Returns:
        成功写入的块数
    """
    try:
        data: list[list[Any]] = [
            [c["chunk_id"] for c in chunks],
            [c["doc_id"] for c in chunks],
            [c["content"] for c in chunks],
            [c["embedding"] for c in chunks],
            [c.get("category", "") for c in chunks],
            [c.get("doc_type", "") for c in chunks],
            [
                ",".join(c.get("keywords", [])) if isinstance(c.get("keywords"), list) else (c.get("keywords", "") or "")
                for c in chunks
            ],
            [c.get("card_type", "") for c in chunks],
            [c.get("customer_tier", "") for c in chunks],
            [int(c.get("effective_date", 0) * 1000) if c.get("effective_date", 0) else 0 for c in chunks],
            [int(c.get("expiry_date", 0) * 1000) if c.get("expiry_date", 0) else 0 for c in chunks],
        ]
        await asyncio.to_thread(collection.insert, data)
        return len(chunks)
    except Exception:
        logger.exception("Milvus 写入失败")
        return 0


# ══════════════════════════════════════════════════════════════
# 6. Publish 阶段
# ══════════════════════════════════════════════════════════════


async def publish_kafka_event(
    doc_id: str,
    chunk_count: int,
    status: str,
    kafka_producer: Any,
    topic: str = "smartcs.knowledge.update",
) -> bool:
    """发布 Kafka 事件通知文档处理完成

    Args:
        doc_id: 文档 ID
        chunk_count: 分块数量
        status: 处理状态
        kafka_producer: aiokafka AIOKafkaProducer
        topic: Kafka topic

    Returns:
        是否发布成功
    """
    try:
        import json

        message = json.dumps(
            {
                "doc_id": doc_id,
                "chunk_count": chunk_count,
                "status": status,
                "timestamp": datetime.now().isoformat(),
            }
        ).encode("utf-8")
        await kafka_producer.send_and_wait(topic, message)
        return True
    except Exception:
        logger.exception("Kafka 发布失败: doc_id=%s", doc_id)
        return False


# ══════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════


async def _log_stage(
    session: AsyncSession,
    doc_id: Any,
    stage: KbIngestionStage,
    status: KbIngestionStatus,
    duration_ms: int,
    step_detail: dict | None = None,
) -> None:
    """记录摄入流水日志

    Args:
        session: 数据库会话
        doc_id: 文档 ID
        stage: 摄入阶段
        status: 阶段状态
        duration_ms: 阶段耗时（毫秒）
        step_detail: 阶段详细信息
    """
    log = KbIngestionLog(
        document_id=doc_id,
        stage=stage,
        status=status,
        duration_ms=duration_ms,
        step_detail=step_detail,
    )
    session.add(log)
    await session.flush()


# ══════════════════════════════════════════════════════════════
# 编排器
# ══════════════════════════════════════════════════════════════


async def ingest_document(
    doc_id: Any,
    file_path: str,
    source_type: KbSourceType,
    metadata: DocumentMetadata,
    embedding_provider: EmbeddingProvider,
    db_session: AsyncSession,
    es_client: Any | None = None,
    milvus_collection: Any | None = None,
    kafka_producer: Any | None = None,
    chunk_size: int = 1500,
    chunk_overlap: int = 200,
) -> str:
    """文档摄入编排器

    6 阶段流水线：Parse → Clean → Chunk → Embed → Dual-Write → Publish

    Args:
        doc_id: 文档 UUID
        file_path: 文件路径（MARKDOWN 类型时为文本内容）
        source_type: 文档来源类型
        metadata: 文档元数据
        embedding_provider: 嵌入服务提供者
        db_session: 数据库异步会话
        es_client: Elasticsearch 异步客户端（可选）
        milvus_collection: Milvus Collection 对象（可选）
        kafka_producer: Kafka 生产者（可选）
        chunk_size: 分块大小
        chunk_overlap: 分块重叠

    Returns:
        最终状态: COMPLETED / PARTIAL_ES_ONLY / KAFKA_PENDING / FAILED
    """
    # 查询文档记录
    doc = await db_session.get(KbDocument, doc_id)
    if doc is None:
        raise DocumentFormatError(f"文档不存在: {doc_id}")

    # 设置为 PROCESSING
    doc.status = KbDocStatus.PROCESSING
    await db_session.flush()

    try:
        # ── 1. Parse ──
        t0 = time.perf_counter()
        raw_text = _parse(source_type, file_path)
        await _log_stage(
            db_session,
            doc_id,
            KbIngestionStage.PARSE,
            KbIngestionStatus.SUCCESS,
            int((time.perf_counter() - t0) * 1000),
        )

        # ── 2. Clean ──
        t0 = time.perf_counter()
        cleaned = clean_text(raw_text)
        await _log_stage(
            db_session,
            doc_id,
            KbIngestionStage.CLEAN,
            KbIngestionStatus.SUCCESS,
            int((time.perf_counter() - t0) * 1000),
        )

        # ── 3. Chunk ──
        t0 = time.perf_counter()
        chunks = chunk_text(cleaned, chunk_size=chunk_size, overlap=chunk_overlap)
        await _log_stage(
            db_session,
            doc_id,
            KbIngestionStage.CHUNK,
            KbIngestionStatus.SUCCESS,
            int((time.perf_counter() - t0) * 1000),
            step_detail={"chunk_count": len(chunks)},
        )

        # ── 4. Embed ──
        t0 = time.perf_counter()
        embeddings = await embed_chunks(chunks, embedding_provider)
        await _log_stage(
            db_session,
            doc_id,
            KbIngestionStage.EMBED,
            KbIngestionStatus.SUCCESS,
            int((time.perf_counter() - t0) * 1000),
            step_detail={"embedding_dim": len(embeddings[0]) if embeddings else 0},
        )

        # ── 构建写入数据 ──
        # 将 effective_date / expiry_date 转为 epoch int64
        eff_epoch = 0
        exp_epoch = 0
        if metadata.effective_date:
            try:
                eff_epoch = int(datetime.fromisoformat(metadata.effective_date).timestamp())
            except (ValueError, TypeError):
                eff_epoch = 0
        if metadata.expiry_date:
            try:
                exp_epoch = int(datetime.fromisoformat(metadata.expiry_date).timestamp())
            except (ValueError, TypeError):
                exp_epoch = 0

        chunk_records: list[dict] = []
        for _idx, (chunk_text_str, embedding) in enumerate(zip(chunks, embeddings, strict=True)):
            chunk_id = str(KbIngestionLog.id.default.arg())  # 使用 UUID v7 生成
            chunk_records.append(
                {
                    "chunk_id": chunk_id,
                    "doc_id": str(doc_id),
                    "content": chunk_text_str,
                    "embedding": embedding,
                    "category": metadata.category,
                    "doc_type": metadata.doc_type,
                    "keywords": metadata.keywords,
                    "card_type": metadata.card_type or "",
                    "customer_tier": metadata.customer_tier or "",
                    "effective_date": eff_epoch,
                    "expiry_date": exp_epoch,
                }
            )

        # ── 5. Dual-Write ──
        # 5a. ES 先写
        es_ok = True
        if es_client is not None:
            t0 = time.perf_counter()
            es_count = await write_to_es(chunk_records, es_client)
            es_ok = es_count == len(chunk_records)
            await _log_stage(
                db_session,
                doc_id,
                KbIngestionStage.ES_WRITE,
                KbIngestionStatus.SUCCESS if es_ok else KbIngestionStatus.FAILED,
                int((time.perf_counter() - t0) * 1000),
                step_detail={"success_count": es_count, "total": len(chunk_records)},
            )

        if not es_ok:
            doc.status = KbDocStatus.FAILED
            doc.chunk_count = 0
            await db_session.flush()
            return "FAILED"

        # 5b. Milvus 后写
        milvus_ok = True
        if milvus_collection is not None:
            t0 = time.perf_counter()
            milvus_count = await write_to_milvus(chunk_records, milvus_collection)
            milvus_ok = milvus_count == len(chunk_records)
            await _log_stage(
                db_session,
                doc_id,
                KbIngestionStage.MILVUS_WRITE,
                KbIngestionStatus.SUCCESS if milvus_ok else KbIngestionStatus.FAILED,
                int((time.perf_counter() - t0) * 1000),
                step_detail={"success_count": milvus_count, "total": len(chunk_records)},
            )

        if not milvus_ok:
            doc.status = KbDocStatus.FAILED
            doc.chunk_count = len(chunks)
            await db_session.flush()
            return "PARTIAL_ES_ONLY"

        # ── 保存 KbChunk 到 DB ──
        for idx, record in enumerate(chunk_records):
            chunk = KbChunk(
                id=record["chunk_id"],
                document_id=doc_id,
                chunk_index=idx,
                content=record["content"],
                char_count=len(record["content"]),
                embedding_status=KbEmbedStatus.COMPLETED,
                es_indexed=es_client is not None,
                milvus_indexed=milvus_collection is not None,
            )
            db_session.add(chunk)
        await db_session.flush()

        # ── 6. Publish ──
        final_status = "COMPLETED"
        if kafka_producer is not None:
            t0 = time.perf_counter()
            kafka_ok = await publish_kafka_event(str(doc_id), len(chunks), "COMPLETED", kafka_producer)
            await _log_stage(
                db_session,
                doc_id,
                KbIngestionStage.KAFKA_PUBLISH,
                KbIngestionStatus.SUCCESS if kafka_ok else KbIngestionStatus.FAILED,
                int((time.perf_counter() - t0) * 1000),
            )
            if not kafka_ok:
                final_status = "KAFKA_PENDING"

        # 更新文档状态
        doc.status = KbDocStatus.COMPLETED if final_status == "COMPLETED" else KbDocStatus.KAFKA_PENDING
        doc.chunk_count = len(chunks)
        await db_session.flush()

        return final_status

    except Exception:
        logger.exception("文档摄入失败: doc_id=%s", doc_id)
        doc.status = KbDocStatus.FAILED
        await db_session.flush()
        return "FAILED"

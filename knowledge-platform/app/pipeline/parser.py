"""文档解析器 — docling 版式感知解析

相比 pymupdf 纯文本提取，docling 保留表格结构、标题层级、列表语义，
对银行文档（表格密集）的检索召回至关重要。

支持格式：PDF / DOCX / HTML / Markdown / TXT / XLSX
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from app.orm.kb import KbSourceType

logger = logging.getLogger(__name__)


def parse_markdown(content: str) -> str:
    """从 Markdown 提取结构化文本

    保留标题层级、表格、列表结构，移除 YAML frontmatter。
    """
    from markdown_it import MarkdownIt

    # 移除 YAML frontmatter
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            content = parts[2].strip()

    md = MarkdownIt()
    tokens = md.parse(content)
    parts: list[str] = []
    for token in tokens:
        if token.type == "inline":
            for child in token.children or []:
                if child.type == "text":
                    parts.append(child.content)
        elif (
            token.type in ("heading_open", "paragraph_open", "bullet_list_open", "ordered_list_open")
            and parts
            and parts[-1] != "\n"
        ):
            parts.append("\n")
    return "\n".join(line.strip() for line in "".join(parts).split("\n") if line.strip())


def parse_html(content: str) -> str:
    """从 HTML 提取正文文本"""
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "lxml")
    for tag in soup.find_all(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _parse_with_docling(file_path: str) -> str:
    """使用 docling 将文档转为结构化 Markdown

    docling 的 DocumentConverter 统一处理 PDF/DOCX/PPTX 等格式，
    保留表格结构、标题层级、列表语义。
    """
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(file_path)
    return result.document.export_to_markdown()


def parse_pdf(file_path: str) -> str:
    """从 PDF 提取版式感知文本

    优先 docling（保留表格语义），回退 pymupdf。
    """
    try:
        return _parse_with_docling(file_path)
    except ImportError:
        logger.warning("docling 未安装，回退到 pymupdf")
    except Exception as e:
        logger.warning("docling 解析 PDF 失败，回退 pymupdf: %s", e)

    import fitz

    doc = fitz.open(file_path)
    parts: list[str] = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts)


def parse_docx(file_path: str) -> str:
    """从 DOCX 提取版式感知文本

    优先 docling（保留表格），回退 python-docx。
    """
    try:
        return _parse_with_docling(file_path)
    except ImportError:
        logger.warning("docling 未安装，回退到 python-docx")
    except Exception as e:
        logger.warning("docling 解析 DOCX 失败，回退 python-docx: %s", e)

    from docx import Document

    doc = Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_xlsx(file_path: str) -> str:
    """从 XLSX 提取结构化文本"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h) if h is not None else "" for h in rows[0]]
        for row in rows[1:]:
            cells: list[str] = []
            for header, value in zip(headers, row, strict=False):
                if value is not None:
                    cells.append(f"{header}: {value}")
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def parse_text_content(content: str) -> str:
    """纯文本直通"""
    return content.strip()


_PARSE_DISPATCH: dict[str, Any] = {
    KbSourceType.MARKDOWN: parse_markdown,
    KbSourceType.HTML: parse_html,
    KbSourceType.PDF: parse_pdf,
    KbSourceType.DOCX: parse_docx,
    KbSourceType.XLSX: parse_xlsx,
    KbSourceType.TXT: parse_text_content,
}


def parse(source_type: KbSourceType | str, file_path_or_content: str) -> str:
    """根据来源类型选择解析器

    Args:
        source_type: 文档来源类型
        file_path_or_content: 文件路径（PDF/DOCX/XLSX）或文本内容（MARKDOWN/HTML/TXT）

    Returns:
        解析后的文本
    """
    st = KbSourceType(source_type) if isinstance(source_type, str) else source_type
    parser_fn = _PARSE_DISPATCH.get(st)
    if parser_fn is None:
        raise ValueError(f"不支持的文档格式: {source_type}")
    return parser_fn(file_path_or_content)


def detect_source_type(filename: str) -> KbSourceType:
    """从文件名推断来源类型"""
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": KbSourceType.PDF,
        ".docx": KbSourceType.DOCX,
        ".html": KbSourceType.HTML,
        ".htm": KbSourceType.HTML,
        ".md": KbSourceType.MARKDOWN,
        ".markdown": KbSourceType.MARKDOWN,
        ".txt": KbSourceType.TXT,
        ".xlsx": KbSourceType.XLSX,
    }
    result = mapping.get(ext)
    if result is None:
        raise ValueError(f"无法识别文件类型: {filename}")
    return result
